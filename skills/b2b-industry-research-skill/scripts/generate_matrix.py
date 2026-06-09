#!/usr/bin/env python3
"""
生成竞品-场景交叉矩阵Word文档
使用python-docx库创建格式化的矩阵表格
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if level == 1:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h


def build_matrix(industry, scenes, competitors, output_path=None):
    """
    构建竞品-场景交叉矩阵
    
    Args:
        industry: 行业名称
        scenes: 场景列表，每个场景为字典 {name, type, pain_points}
        competitors: 竞品列表，每个竞品为字典 {name, camp, coverage}
        output_path: 输出文件路径
    
    Returns:
        Document对象
    """
    doc = Document()
    
    # 全局样式
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'{industry}行业竞品场景交叉矩阵表')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_paragraph()
    
    # 按场景类型分组
    scene_types = {}
    for scene in scenes:
        t = scene.get('type', '其他')
        if t not in scene_types:
            scene_types[t] = []
        scene_types[t].append(scene)
    
    # 为每种场景类型生成一个矩阵表
    for scene_type, type_scenes in scene_types.items():
        add_heading_styled(doc, f'{scene_type}场景矩阵', 2)
        
        # 创建表格
        # 列数 = 场景名称列 + 竞品数量
        num_cols = 1 + len(competitors)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置列宽
        table.columns[0].width = Cm(4)
        for i in range(1, num_cols):
            table.columns[i].width = Cm(3)
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '场景名称'
        set_cell_shading(header_cells[0], '1F4E79')
        
        for i, comp in enumerate(competitors):
            header_cells[i+1].text = comp['name']
            set_cell_shading(header_cells[i+1], '2E75B6')
        
        # 设置表头文字样式
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(9)
        
        # 数据行
        for scene in type_scenes:
            row = table.add_row()
            cells = row.cells
            
            # 场景名称
            cells[0].text = scene['name']
            
            # 每个竞品的覆盖情况
            for i, comp in enumerate(competitors):
                coverage = comp.get('coverage', {}).get(scene['name'], '')
                cells[i+1].text = coverage
                
                # 根据市场占有率设置颜色
                market_share = comp.get('market_share', {}).get(scene['name'], '')
                if market_share == 'high':
                    set_cell_shading(cells[i+1], 'FDE8E8')  # 红色背景
                elif market_share == 'medium':
                    set_cell_shading(cells[i+1], 'FFF3E0')  # 橙色背景
                elif market_share == 'low':
                    set_cell_shading(cells[i+1], 'E8F0FE')  # 蓝色背景
            
            # 设置文字样式
            for cell in cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        
        doc.add_paragraph()
    
    # 图例说明
    legend = doc.add_paragraph()
    legend.add_run('图例说明：').font.bold = True
    doc.add_paragraph('红色背景：市场占有率极高（70-90%）', style='List Bullet')
    doc.add_paragraph('橙色背景：市场占有率中等', style='List Bullet')
    doc.add_paragraph('蓝色背景：市场占有率低', style='List Bullet')
    
    if output_path:
        doc.save(output_path)
    
    return doc


if __name__ == '__main__':
    # 示例用法
    industry = "电力"
    
    scenes = [
        {"name": "两票三制管理", "type": "安全生产", "pain_points": "纸质票据效率低、审批流程长"},
        {"name": "智慧工地", "type": "建设管理", "pain_points": "安全管理粗放、物资损耗大"},
        {"name": "虚拟电厂", "type": "能源运营", "pain_points": "资源聚合难、调度响应慢"},
    ]
    
    competitors = [
        {"name": "国网信通", "camp": "国网系", "coverage": {"两票三制管理": "全面支持", "智慧工地": "部分支持", "虚拟电厂": "全面支持"}, "market_share": {"两票三制管理": "high", "虚拟电厂": "high"}},
        {"name": "南瑞集团", "camp": "国网系", "coverage": {"两票三制管理": "全面支持", "智慧工地": "不支持", "虚拟电厂": "全面支持"}, "market_share": {"两票三制管理": "high", "虚拟电厂": "medium"}},
        {"name": "远光软件", "camp": "第三方", "coverage": {"两票三制管理": "部分支持", "智慧工地": "不支持", "虚拟电厂": "部分支持"}, "market_share": {"虚拟电厂": "low"}},
    ]
    
    doc = build_matrix(industry, scenes, competitors, "示例矩阵.docx")
    print("矩阵文档已生成：示例矩阵.docx")
